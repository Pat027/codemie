# Copyright 2026 EPAM Systems, Inc. ("EPAM")
#
# Licensed under the Apache License, Version 2.0 (the "License");
# you may not use this file except in compliance with the License.
# You may obtain a copy of the License at
#
#     http://www.apache.org/licenses/LICENSE-2.0
#
# Unless required by applicable law or agreed to in writing, software
# distributed under the License is distributed on an "AS IS" BASIS,
# WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
# See the License for the specific language governing permissions and
# limitations under the License.

"""DOCX workers for multiprocessing."""

import io
import logging
from typing import Any
from codemie_tools.file_analysis.docx.exceptions import TableExtractionError
from docx import Document
from codemie_tools.file_analysis.docx.models import (
    HeaderInfo,
    ParagraphInfo,
    Position,
    StyleInfo,
    SectionInfo,
    DocumentContent,
    DocumentStructure,
    FormattingInfo,
    TableData,
)

logger = logging.getLogger(__name__)


def extract_docx_text(docx_bytes: bytes) -> str:
    """
    Extract text from DOCX document.

    Args:
        docx_bytes: DOCX file content as bytes

    Returns:
        Extracted text content
    """
    try:
        document = Document(io.BytesIO(docx_bytes))
        return _extract_text(document)
    except Exception as e:
        logger.error(f"DOCX text extraction failed: {e}")
        raise


def extract_docx_content(
    docx_bytes: bytes,
    /,
    extract_text: bool = True,
    extract_structure: bool = False,
    extract_formatting: bool = False,
    extract_metadata: bool = False,
    extract_tables: bool = False,
    *_: Any,
    **__: Any,
) -> Any:
    """
    Extract content from DOCX document in subprocess.

    Args:
        docx_bytes: DOCX file content as bytes
        extract_text: Whether to extract plain text
        extract_structure: Whether to extract document structure
        extract_formatting: Whether to extract formatting info
        extract_metadata: Whether to extract document metadata
        extract_tables: Whether to extract tables

    Returns:
        DocumentContent object
    """
    try:
        document = Document(io.BytesIO(docx_bytes))

        text = _extract_text(document) if extract_text else ""
        structure = _extract_structure(document) if extract_structure else DocumentStructure()
        formatting = _extract_formatting(document) if extract_formatting else FormattingInfo()
        metadata = _extract_metadata(document) if extract_metadata else {}
        tables = _extract_tables(document) if extract_tables else []

        return DocumentContent(
            text=text,
            structure=structure,
            formatting=formatting,
            metadata=metadata,
            tables=tables,
            images=[],
        )
    except Exception as e:
        logger.error(f"DOCX content extraction failed: {e}")
        raise


def _extract_text(document) -> str:
    """
    Extract text content from a document.

    Args:
        document: python-docx Document object

    Returns:
        Document text content
    """

    text_parts = []
    for paragraph in document.paragraphs:
        text_parts.append(paragraph.text)
    for table in document.tables:
        for row in table.rows:
            row_text = [cell.text for cell in row.cells]
            text_parts.append(" | ".join(row_text))
    return "\n".join(text_parts)


def _extract_structure(document) -> Any:
    paragraphs, headers, styles = _process_paragraphs(document.paragraphs)
    sections = _create_sections(paragraphs, headers)
    return DocumentStructure(headers=headers, paragraphs=paragraphs, sections=sections, styles=styles)


def _process_paragraphs(doc_paragraphs) -> tuple:
    """
    Extract document structure information.

    Args:
        document: python-docx Document object

    Returns:
        DocumentStructure object
    """

    headers = []
    paragraphs = []
    styles = []

    for i, paragraph in enumerate(doc_paragraphs):
        position = Position(page=i // 40 + 1, x=0.0, y=float(i % 40))
        style_name = paragraph.style.name if paragraph.style and paragraph.style.name else "Normal"

        if style_name.startswith("Heading"):
            level = int(style_name.replace("Heading ", "")) if style_name != "Heading" else 1
            headers.append(HeaderInfo(level=level, text=paragraph.text, position=position))

        paragraphs.append(ParagraphInfo(text=paragraph.text, style=style_name, position=position))
        _add_style_if_new(paragraph, styles, style_name, StyleInfo)

    return paragraphs, headers, styles


def _add_style_if_new(paragraph, styles: list, default_style_name: str, style_info) -> None:  # ignore
    """
    Add style information if not already processed.

    Args:
        paragraph: Document paragraph
        styles: List of styles to update
    """
    style_name = paragraph.style.name if paragraph.style and paragraph.style.name else default_style_name
    if style_name not in [s.name for s in styles]:
        styles.append(
            style_info(
                name=style_name,
                font=paragraph.style.font.name
                if hasattr(paragraph.style, "font") and paragraph.style.font
                else "Default",
                size=paragraph.style.font.size if hasattr(paragraph.style, "font") and paragraph.style.font else 12,
                bold=paragraph.bold if hasattr(paragraph, "bold") else False,
                italic=paragraph.italic if hasattr(paragraph, "italic") else False,
            )
        )


def _create_sections(paragraphs: list, headers: list) -> list:
    """
    Create document sections based on headers.

    Args:
        paragraphs: List of paragraph info objects
        headers: List of header info objects

    Returns:
        List of section info objects
    """

    sections = []
    current_section = None
    section_content: list = []

    for para in paragraphs:
        # Check if this paragraph is a header
        header_match = next((h for h in headers if h.text == para.text), None)

        if header_match:
            _finalize_section(current_section, section_content, sections)
            current_section = header_match
            section_content = []
        else:
            # Add to current section content
            section_content.append(para)

    _finalize_section(current_section, section_content, sections)
    return sections


def _finalize_section(current_section, section_content, sections):
    """
    Add a section to the sections list if it exists.

    Args:
        current_section: Current section header
        section_content: Content of the section
        sections: List of sections to update
    """
    if current_section:
        sections.append(
            SectionInfo(
                title=current_section.text,
                content=section_content.copy(),
                level=current_section.level,
            )
        )


def _extract_formatting(document) -> Any:
    """
    Extract formatting information from a document.

    Args:
        document: python-docx Document object

    Returns:
        FormattingInfo object
    """

    style_dict = {}
    for style in document.styles:
        if hasattr(style, "name") and hasattr(style, "font"):
            style_dict[style.name] = StyleInfo(
                name=style.name,
                font=style.font.name if hasattr(style.font, "name") else "Default",
                size=style.font.size if hasattr(style.font, "size") else 12,
            )

    page_width, page_height, margins = _extract_page_dimensions(document)
    return FormattingInfo(styles=style_dict, page_width=page_width, page_height=page_height, margins=margins)


def _extract_page_dimensions(document) -> tuple[float, float, dict]:
    """
    Extract page dimensions from document.

    Args:
        document: python-docx Document object

    Returns:
        Tuple of (page_width, page_height, margins)
    """
    default_width = 8.5
    default_height = 11.0
    default_margins: dict[str, float] = {"top": 1.0, "right": 1.0, "bottom": 1.0, "left": 1.0}

    try:
        section = document.sections[0]
        page_width = section.page_width.inches if hasattr(section, "page_width") else default_width
        page_height = section.page_height.inches if hasattr(section, "page_height") else default_height
        margins = {
            "top": section.top_margin.inches if hasattr(section, "top_margin") else default_margins["top"],
            "right": section.right_margin.inches if hasattr(section, "right_margin") else default_margins["right"],
            "bottom": section.bottom_margin.inches if hasattr(section, "bottom_margin") else default_margins["bottom"],
            "left": section.left_margin.inches if hasattr(section, "left_margin") else default_margins["left"],
        }
        return page_width, page_height, margins
    except (AttributeError, IndexError):
        return default_width, default_height, default_margins


def _extract_metadata(document) -> dict[str, Any]:
    """
    Extract metadata from a document.

    Args:
        document: python-docx Document object

    Returns:
        Dictionary of metadata
    """
    metadata: dict[str, Any] = {}
    core_props = document.core_properties

    for attr in (
        "title",
        "author",
        "keywords",
        "subject",
        "category",
    ):
        if hasattr(core_props, attr):
            metadata[attr] = getattr(core_props, attr)

    if hasattr(core_props, "created"):
        metadata["created_date"] = core_props.created.isoformat() if core_props.created else None
    if hasattr(core_props, "modified"):
        metadata["modified_date"] = core_props.modified.isoformat() if core_props.modified else None

    metadata["paragraph_count"] = len(document.paragraphs)
    metadata["table_count"] = len(document.tables)
    metadata["section_count"] = len(document.sections)
    metadata["word_count"] = sum(len(p.text.split()) for p in document.paragraphs)
    return metadata


def _extract_tables(document) -> list:
    """
    Extract tables from a document.

    Args:
        document: python-docx Document object

    Returns:
        List of TableData objects

    Raises:
        TableExtractionError: If table extraction fails
    """

    try:
        tables = []
        for i, table in enumerate(document.tables):
            rows = []
            headers: list[str] = []
            for j, row in enumerate(table.rows):
                row_data = [cell.text for cell in row.cells]
                if j == 0:
                    headers = row_data
                rows.append(row_data)
            tables.append(
                TableData(
                    rows=rows,
                    headers=headers,
                    position=Position(page=i + 1, x=0.0, y=0.0),
                    metadata={"table_index": i},
                )
            )
    except Exception as e:
        raise TableExtractionError(f"Failed to extract tables: {str(e)}") from e
    return tables
