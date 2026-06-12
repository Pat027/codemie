# Copyright 2026 EPAM Systems, Inc. ("EPAM")
#
# Licensed under the Apache License, Version 2.0 (the "License");
# you may not use this file except in compliance with the License.
# You may obtain a copy of the License at
#
#     http://www.apache.org/licenses/LICENSE-2.0
#
# Unless required by applicable law or agreed to in writing, software
# distributed under the License is distributed on an "AS IS" BASIS,
# WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
# See the License for the specific language governing permissions and
# limitations under the License.

"""Unit tests for DocxLoader (markitdown with python-docx fallback)."""

import os
import tempfile
from unittest.mock import MagicMock, patch

from docx import Document as DocxDocument
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from langchain_core.documents import Document

from codemie.datasource.loader.docx_loader import DocxLoader


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _save_tmp(doc) -> str:
    fd, path = tempfile.mkstemp(suffix=".docx")
    os.close(fd)
    doc.save(path)
    return path


def _build_normal_docx() -> str:
    doc = DocxDocument()
    doc.add_heading("Quarterly Report", level=1)
    doc.add_paragraph("Some intro text.")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Name"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "Alpha"
    table.cell(1, 1).text = "42"
    return _save_tmp(doc)


def _build_broken_docx() -> str:
    """A DOCX whose table row contains a bookmark — crashes mammoth/markitdown.

    Reproduces ``AttributeError: 'TableCellUnmerged' object has no attribute '_accept1'``.
    """
    doc = DocxDocument()
    doc.add_paragraph("Intro paragraph.")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "h1"
    table.cell(0, 1).text = "h2"
    table.cell(1, 0).text = "v1"
    table.cell(1, 1).text = "v2"

    tr = table.rows[0]._tr
    bookmark = OxmlElement("w:bookmarkStart")
    bookmark.set(qn("w:id"), "0")
    bookmark.set(qn("w:name"), "row_level_bookmark")
    tr.insert(0, bookmark)
    return _save_tmp(doc)


# ---------------------------------------------------------------------------
# Happy path: markitdown is used for files it can convert
# ---------------------------------------------------------------------------


def test_normal_docx_uses_markitdown_not_fallback():
    path = _build_normal_docx()
    try:
        docs = list(DocxLoader(path).lazy_load())
        assert len(docs) == 1
        assert isinstance(docs[0], Document)
        # markitdown path does not set the fallback marker
        assert "conversion_fallback" not in docs[0].metadata
        assert "Quarterly Report" in docs[0].page_content
    finally:
        os.unlink(path)


# ---------------------------------------------------------------------------
# Fallback path: markitdown fails -> python-docx extraction
# ---------------------------------------------------------------------------


def test_broken_docx_falls_back_to_python_docx():
    path = _build_broken_docx()
    try:
        docs = list(DocxLoader(path).lazy_load())
        assert len(docs) == 1
        assert docs[0].metadata["conversion_fallback"] == "python-docx"
        assert docs[0].metadata["conversion_success"] is True
    finally:
        os.unlink(path)


def test_fallback_extracts_paragraph_and_table_text():
    path = _build_broken_docx()
    try:
        content = list(DocxLoader(path).lazy_load())[0].page_content
        assert "Intro paragraph." in content
        # table rendered as a markdown pipe table
        assert "| h1 | h2 |" in content
        assert "| v1 | v2 |" in content
        assert "| --- | --- |" in content
    finally:
        os.unlink(path)


def test_fallback_renders_heading_as_markdown():
    doc = DocxDocument()
    doc.add_heading("Section Title", level=2)
    doc.add_paragraph("Body text.")
    # Force the fallback by inserting a row-level bookmark in a table
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "cell"
    bm = OxmlElement("w:bookmarkStart")
    bm.set(qn("w:id"), "0")
    bm.set(qn("w:name"), "bm")
    table.rows[0]._tr.insert(0, bm)
    path = _save_tmp(doc)
    try:
        content = list(DocxLoader(path).lazy_load())[0].page_content
        assert "## Section Title" in content
        assert "Body text." in content
    finally:
        os.unlink(path)


def test_yields_nothing_when_markitdown_and_python_docx_both_fail():
    loader = DocxLoader("/nonexistent/file.docx")
    loader._markitdown_loader = MagicMock()
    loader._markitdown_loader.load.side_effect = ValueError("markitdown failed")
    with patch("docx.Document", side_effect=Exception("python-docx failed")):
        docs = list(loader.lazy_load())
    assert docs == []


# ---------------------------------------------------------------------------
# Rendering helpers
# ---------------------------------------------------------------------------


def test_render_table_single_row_has_no_separator():
    doc = DocxDocument()
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "a"
    table.cell(0, 1).text = "b"
    rendered = DocxLoader._render_table(table)
    assert rendered == "| a | b |"
    assert "---" not in rendered


def test_render_table_multi_row_has_header_separator():
    doc = DocxDocument()
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "h1"
    table.cell(0, 1).text = "h2"
    table.cell(1, 0).text = "v1"
    table.cell(1, 1).text = "v2"
    rendered = DocxLoader._render_table(table)
    assert rendered.splitlines()[1] == "| --- | --- |"


def test_render_table_pipe_in_cell_is_escaped():
    doc = DocxDocument()
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "yes|no"
    table.cell(0, 1).text = "col2"
    table.cell(1, 0).text = "a"
    table.cell(1, 1).text = "b"
    rendered = DocxLoader._render_table(table)
    lines = rendered.splitlines()
    # pipe in cell text is escaped; separator has exactly 2 columns
    assert r"yes\|no" in lines[0]
    assert lines[1] == "| --- | --- |"


def test_fallback_empty_docx_yields_nothing():
    doc = DocxDocument()
    # force fallback by inserting a row-level bookmark in a table
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = ""
    bm = OxmlElement("w:bookmarkStart")
    bm.set(qn("w:id"), "0")
    bm.set(qn("w:name"), "bm")
    table.rows[0]._tr.insert(0, bm)
    path = _save_tmp(doc)
    try:
        docs = list(DocxLoader(path).lazy_load())
        assert docs == []
    finally:
        os.unlink(path)
